from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx_utils import set_paragraph_spacing, add_hyperlink, add_horizontal_line
import pypandoc
import os
from generate_pdfs import save_as_pdf


def generate_compact_resume(data, output_file="compact_resume.docx",generate_pdf=False):
    """
    Generate a compact Word document resume.
    """
    doc = Document()

    # Header Section
    name = doc.add_heading(level=1)
    name_run = name.add_run(data["header"]["name"])
    name_run.bold = True
    name_run.font.size = Pt(16)
    name.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    set_paragraph_spacing(name, after=2)

    header_paragraph = doc.add_paragraph()
    header_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    header_paragraph.add_run("P: ").bold = True
    header_paragraph.add_run(data["header"]["contact"]["phone"] + "    ")

    # Add LinkedIn with hyperlink
    add_hyperlink(header_paragraph, "LinkedIn", data["header"]["contact"]["linkedin"])
    header_paragraph.add_run("    ")

    # Add Email with hyperlink
    add_hyperlink(header_paragraph, "Email", f"mailto:{data['header']['contact']['email']}")
    set_paragraph_spacing(header_paragraph, after=2)

    # Add horizontal line below header
    add_horizontal_line(header_paragraph)

    # Summary Section
    doc.add_heading("Summary", level=2)
    summary_paragraph = doc.add_paragraph(data["summary"])
    summary_paragraph.style.font.size = Pt(11)
    set_paragraph_spacing(summary_paragraph, after=6)
    # Add horizontal line below header
    add_horizontal_line(summary_paragraph)
    
    # Education Section
    doc.add_heading("Education", level=2)
    for edu in data["education"]:
        education_paragraph = doc.add_paragraph()
        education_paragraph.add_run(f"{edu['degree']} - {edu['university']} | ").bold = True
        education_paragraph.add_run(f"GPA: {edu.get('gpa', 'N/A')} ").italic = True
        education_paragraph.add_run(f"({edu.get('graduation_date', 'N/A')})")
        set_paragraph_spacing(education_paragraph, after=2)
    # Add horizontal line below header
    add_horizontal_line(education_paragraph)
    
    
    # Technical Skills Section
    doc.add_heading("Technical Skills", level=2)
    # Access the nested dictionary under the "Skills" key
    skills_data = data["skills"].get("Skills", {})  # Get the nested dictionary or an empty one if missing
    if not skills_data:
        print("No skills data found. Skipping Technical Skills section.")
    else:
        # Add each category and its skills in a single line
        for category, skills in skills_data.items():
        # Create a new paragraph for each category and skills
            skills_paragraph = doc.add_paragraph()
            # Add the category in bold
            category_run = skills_paragraph.add_run(f"{category}: ")
            category_run.bold = True
            category_run.font.size = Pt(10)
            # Add the skills in regular font
            skills_run = skills_paragraph.add_run(", ".join(skills))
            skills_run.font.size = Pt(10)
            # Adjust spacing
            set_paragraph_spacing(skills_paragraph, after=1)

        # Add horizontal line below header
        add_horizontal_line(skills_paragraph)

    # Work Experience Section
    doc.add_heading("Work Experience", level=2)
    for exp in data["experience"]:
        work_heading = doc.add_paragraph()
        work_heading.add_run(f"{exp['title']} - {exp['company']} | ").bold = True
        work_heading.add_run(f"({exp['duration']})").italic = True
        set_paragraph_spacing(work_heading, after=2)
        for responsibility in exp.get("responsibilities", []):
            responsibility_paragraph = doc.add_paragraph(f"{responsibility}", style="List Bullet")
            responsibility_paragraph.style.font.size = Pt(10)
            set_paragraph_spacing(responsibility_paragraph, after=1)
    # Add horizontal line below header
    add_horizontal_line(responsibility_paragraph)
    
    
    # Project Experience Section
    doc.add_heading("Project Experience", level=2)
    for project in data["projects"]:
        # Create a paragraph for the project
        project_paragraph = doc.add_paragraph()

        # Add the project name in bold
        project_name_run = project_paragraph.add_run(f"{project['name']}: ")
        project_name_run.bold = True
        project_name_run.font.size = Pt(11)

        # Add the project description in regular font
        project_description_run = project_paragraph.add_run(project["description"])
        project_description_run.font.size = Pt(10)

        # Adjust spacing
        set_paragraph_spacing(project_paragraph, after=2)

    # Save the document
    doc.save(output_file)
    print(f"Resume generated: {output_file}")
    # Convert to PDF if required
    if generate_pdf:
        save_as_pdf(data, pdf_file=os.path.splitext(output_file)[0] + ".pdf")

